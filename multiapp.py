import tkinter as tk
from tkinter import ttk, filedialog, colorchooser, messagebox, simpledialog
import pygame
import os
import cv2
from PIL import Image, ImageTk, ImageGrab
from docx import Document
import tempfile
import subprocess
import re

class MultiToolApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Мультифункциональное приложение v6.1")
        self.root.geometry("1100x780")
        self.root.minsize(850, 550)
        self.is_running = True  # Флаг для безопасной остановки фоновых задач

        pygame.mixer.init()
        self.current_theme = "light"
        
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.setup_theme_manager()
        self.setup_text_editor()
        self.setup_paint()
        self.setup_media_player()
        self.setup_settings()
        self.setup_global_hotkeys()

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    # ==================== ТЕМЫ ====================
    def setup_theme_manager(self):
        self.themes = {
            "light": {"bg": "#ffffff", "fg": "#000000", "entry_bg": "#ffffff", "btn_bg": "#f0f0f0", 
                      "canvas_bg": "#ffffff", "text_bg": "#ffffff", "hl_bg": "#f5f5f5"},
            "dark": {"bg": "#2b2b2b", "fg": "#d4d4d4", "entry_bg": "#3c3f41", "btn_bg": "#4c5054",
                     "canvas_bg": "#1e1e1e", "text_bg": "#1e1e1e", "hl_bg": "#252526"}
        }
        self.style = ttk.Style()
        self.apply_theme("light")

    def apply_theme(self, theme_name):
        self.current_theme = theme_name
        t = self.themes[theme_name]
        self.style.theme_use("default")
        
        self.style.configure(".", background=t["bg"], foreground=t["fg"])
        self.style.configure("TFrame", background=t["bg"])
        self.style.configure("TLabel", background=t["bg"], foreground=t["fg"])
        self.style.configure("TButton", background=t["btn_bg"], foreground=t["fg"])
        self.style.map("TButton", background=[("active", "#5a9cf8"), ("pressed", "#4070c0")])
        self.style.configure("TScale", background=t["bg"], troughcolor="#666", sliderlength=15)
        self.style.configure("Horizontal.TScale", background=t["bg"])
        self.style.configure("TNotebook", background=t["bg"], borderwidth=0)
        self.style.configure("TNotebook.Tab", padding=[8, 4], background=t["hl_bg"])
        self.style.configure("TSeparator", background="#888888")
        self.style.configure("TLabelframe", background=t["bg"], foreground=t["fg"])
        self.style.configure("TLabelframe.Label", background=t["bg"], foreground=t["fg"])
        
        self.root.config(bg=t["bg"])
        if hasattr(self, "text"):
            self.text.config(bg=t["text_bg"], fg=t["fg"], insertbackground=t["fg"], selectbackground="#0078d7")
        if hasattr(self, "canvas"):
            self.canvas.config(bg=t["canvas_bg"])
        if hasattr(self, "video_label"):
            self.video_label.config(bg="#000000")
        if hasattr(self, "playlist_listbox"):
            self.playlist_listbox.config(bg=t["entry_bg"], fg=t["fg"], selectbackground="#0078d7", selectforeground="#ffffff")
        if hasattr(self, "doc_text"):
            self.doc_text.config(bg=t["text_bg"], fg=t["fg"])

    def toggle_theme(self):
        self.apply_theme("dark" if self.current_theme == "light" else "light")

    # ==================== ТЕКСТОВЫЙ РЕДАКТОР ====================
    def setup_text_editor(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="📝 Редактор")

        self.text = tk.Text(frame, wrap=tk.WORD, font=("DejaVu Sans Mono", 12), undo=True)
        self.text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.highlight_syntax = tk.BooleanVar(value=False)
        self.syntax_timer = None

        toolbar = ttk.Frame(frame)
        toolbar.pack(fill=tk.X, padx=5, pady=2)

        ttk.Button(toolbar, text="📄 Новый", command=self.new_file).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="📂 Открыть", command=self.open_file).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="💾 Сохранить", command=self.save_file).pack(side=tk.LEFT, padx=2)
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)

        self.fonts = ["DejaVu Sans Mono", "Ubuntu Mono", "Consolas", "Courier New", "Arial"]
        self.sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 24]
        self.font_combo = ttk.Combobox(toolbar, values=self.fonts, width=16, state="readonly")
        self.font_combo.current(0)
        self.font_combo.pack(side=tk.LEFT, padx=2)
        self.size_combo = ttk.Combobox(toolbar, values=self.sizes, width=3, state="readonly")
        self.size_combo.current(4)
        self.size_combo.pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="🔤 Шрифт", command=self.apply_font).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="📊 Таблица", command=self.insert_table).pack(side=tk.LEFT, padx=2)
        ttk.Checkbutton(toolbar, text="🎨 Подсветка", variable=self.highlight_syntax, command=self.toggle_highlighting).pack(side=tk.LEFT, padx=10)
        ttk.Button(toolbar, text="🌗 Тема", command=self.toggle_theme).pack(side=tk.RIGHT, padx=2)

        self.current_file = None
        self.status_var = tk.StringVar(value="Ctrl+N/O/S | F1/F2 Тема | Ctrl+Z/Y Отмена/Повтор")
        ttk.Label(frame, textvariable=self.status_var, anchor=tk.W).pack(fill=tk.X, padx=5, pady=2)
        
        self.text.bind("<KeyRelease>", self.schedule_highlight)
        self._setup_syntax_tags()

    def _setup_syntax_tags(self):
        self.text.tag_configure("keyword", foreground="#569cd6", font=("DejaVu Sans Mono", 12, "bold"))
        self.text.tag_configure("string", foreground="#ce9178")
        self.text.tag_configure("comment", foreground="#6a9955", font=("DejaVu Sans Mono", 12, "italic"))
        self.text.tag_configure("number", foreground="#b5cea8")

    def schedule_highlight(self, event=None):
        if self.syntax_timer: self.root.after_cancel(self.syntax_timer)
        if self.highlight_syntax.get(): self.syntax_timer = self.root.after(300, self.highlight_text)

    def highlight_text(self):
        if not self.highlight_syntax.get(): return
        txt = self.text.get("1.0", "end-1c")
        for tag in ("keyword", "string", "comment", "number"):
            self.text.tag_remove(tag, "1.0", tk.END)

        keywords = ["def", "class", "import", "from", "if", "else", "for", "while", "return", 
                    "try", "except", "with", "as", "True", "False", "None", "and", "or", "not"]
        for line_num, line in enumerate(txt.split("\n"), 1):
            for m in re.finditer(r"(#.*)", line):
                self.text.tag_add("comment", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")
            for m in re.finditer(r"(['\"].*?['\"])", line):
                self.text.tag_add("string", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")
            for m in re.finditer(r"\b(\d+\.?\d*)\b", line):
                self.text.tag_add("number", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")
            for kw in keywords:
                for m in re.finditer(rf"\b{kw}\b", line):
                    self.text.tag_add("keyword", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")

    def toggle_highlighting(self):
        if not self.highlight_syntax.get():
            for tag in ("keyword", "string", "comment", "number"):
                self.text.tag_remove(tag, "1.0", tk.END)

    def new_file(self):
        self.text.delete("1.0", tk.END)
        self.current_file = None

    def open_file(self):
        path = filedialog.askopenfilename(filetypes=[("Text/Code/Word", "*.txt *.md *.log *.py *.js *.html *.css *.docx")])
        if not path: return
        try:
            self.text.delete("1.0", tk.END)
            # Попытка чтения с разными кодировками
            content = ""
            for enc in ["utf-8", "utf-8-sig", "cp1251"]:
                try:
                    with open(path, "r", encoding=enc) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            if path.endswith(".docx"):
                doc = Document(path)
                content = "\n".join([p.text for p in doc.paragraphs])
            
            self.text.insert(tk.END, content)
            self.current_file = path
            self.status_var.set(f"Открыт: {os.path.basename(path)}")
            self.text.see("1.0")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{e}")

    def save_file(self):
        if self.current_file:
            path = self.current_file
        else:
            path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                filetypes=[("Text files", "*.txt"), ("Word Document", "*.docx"), ("All files", "*.*")])
        if not path: return
        try:
            if path.endswith(".docx"):
                doc = Document()
                for line in self.text.get("1.0", "end-1c").split("\n"):
                    doc.add_paragraph(line)
                doc.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.text.get("1.0", tk.END))
            self.current_file = path
            self.status_var.set(f"Сохранено: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")

    def apply_font(self):
        try:
            sel = self.text.tag_ranges("sel")
            if not sel:
                messagebox.showinfo("Информация", "Выделите текст, чтобы применить шрифт.")
                return
            font = self.font_combo.get()
            size = self.size_combo.get()
            self.text.tag_configure("custom_font", font=(font, size))
            self.text.tag_add("custom_font", sel[0], sel[1])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось применить шрифт:\n{e}")

    def insert_table(self):
        rows = simpledialog.askinteger("Таблица", "Строки:", minvalue=1, maxvalue=10, initialvalue=3)
        cols = simpledialog.askinteger("Таблица", "Столбцы:", minvalue=1, maxvalue=6, initialvalue=3)
        if rows and cols:
            table_frame = ttk.Frame(self.text, relief=tk.SOLID, borderwidth=1)
            for r in range(rows):
                for c in range(cols):
                    ttk.Entry(table_frame, width=12).grid(row=r, column=c, padx=1, pady=1)
            self.text.window_create(tk.INSERT, window=table_frame)
            self.text.insert(tk.INSERT, "\n\n")
            self.text.see(tk.INSERT)

    # ==================== PAINT ====================
    def setup_paint(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="🎨 Paint")

        toolbar = ttk.Frame(frame)
        toolbar.pack(fill=tk.X, padx=5, pady=2)

        self.brush_color = tk.StringVar(value="black")
        ttk.Button(toolbar, text="🌈 Цвет", command=self.choose_color).pack(side=tk.LEFT, padx=2)
        self.brush_size = tk.IntVar(value=3)
        ttk.Label(toolbar, text="Толщина:").pack(side=tk.LEFT, padx=2)
        ttk.Spinbox(toolbar, from_=1, to=30, textvariable=self.brush_size, width=3).pack(side=tk.LEFT, padx=2)
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)

        tools = {"pencil": "✏️ Карандаш", "line": "📏 Линия", "rect": "⬜ Прямоугольник", "oval": "⭕ Овал", "eraser": "🧹 Ластик"}
        self.tool_combo = ttk.Combobox(toolbar, values=list(tools.values()), state="readonly", width=14)
        self.tool_combo.current(0)
        self.tool_combo.pack(side=tk.LEFT, padx=2)
        self.tool_map = {v: k for k, v in tools.items()}
        
        ttk.Button(toolbar, text="💾 PNG", command=self.save_png).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="💾 SVG", command=self.save_svg).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="🗑 Очистить", command=self.clear_canvas).pack(side=tk.LEFT, padx=2)

        ttk.Separator(frame, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=5)

        canvas_wrapper = ttk.Frame(frame, relief="solid", borderwidth=2, padding=5)
        canvas_wrapper.pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5))

        self.canvas = tk.Canvas(canvas_wrapper, bg="white", cursor="cross", highlightthickness=0)
        self.canvas.pack(fill=tk.BOTH, expand=True)

        self.last_x, self.last_y = None, None
        self.preview_id = None
        self.start_x, self.start_y = None, None

        self.canvas.bind("<Button-1>", self.start_draw)
        self.canvas.bind("<B1-Motion>", self.draw)
        self.canvas.bind("<ButtonRelease-1>", self.end_draw)

    def choose_color(self):
        color = colorchooser.askcolor(title="Цвет кисти")
        if color[1]: self.brush_color.set(color[1])

    def clear_canvas(self):
        self.canvas.delete("all")
        self.preview_id = None

    def start_draw(self, event):
        self.last_x, self.last_y = event.x, event.y
        self.start_x, self.start_y = event.x, event.y
        tool = self.tool_map[self.tool_combo.get()]
        if tool not in ("pencil", "eraser"): return
        color = "white" if tool == "eraser" else self.brush_color.get()
        r = self.brush_size.get()
        self.canvas.create_oval(event.x-r, event.y-r, event.x+r, event.y+r, fill=color, outline=color)

    def draw(self, event):
        tool = self.tool_map[self.tool_combo.get()]
        color = "white" if tool == "eraser" else self.brush_color.get()
        size = self.brush_size.get()

        if tool in ("pencil", "eraser"):
            self.canvas.create_line(self.last_x, self.last_y, event.x, event.y,
                                    fill=color, width=size, capstyle=tk.ROUND, smooth=True)
        else:
            if self.preview_id: self.canvas.delete(self.preview_id)
            x1, y1, x2, y2 = self.start_x, self.start_y, event.x, event.y
            if tool == "line":
                self.preview_id = self.canvas.create_line(x1, y1, x2, y2, fill=color, width=size)
            elif tool == "rect":
                self.preview_id = self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=size)
            elif tool == "oval":
                self.preview_id = self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=size)
            return
        self.last_x, self.last_y = event.x, event.y

    def end_draw(self, event):
        self.last_x, self.last_y = None, None
        self.preview_id = None
        self.start_x, self.start_y = None, None

    def save_png(self):
        path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG", "*.png"), ("JPEG", "*.jpg")])
        if not path: return
        try:
            x = self.root.winfo_rootx() + self.canvas.winfo_x() + 2
            y = self.root.winfo_rooty() + self.canvas.winfo_y() + 2
            x1, y1 = x + self.canvas.winfo_width(), y + self.canvas.winfo_height()
            ImageGrab.grab(bbox=(x, y, x1, y1)).save(path)
            messagebox.showinfo("Успех", f"Сохранено в {path}")
        except Exception as e:
            messagebox.showwarning("Ошибка", f"Не удалось сохранить PNG: {e}")

    def save_svg(self):
        path = filedialog.asksaveasfilename(defaultextension=".svg", filetypes=[("SVG Vector", "*.svg")])
        if not path: return
        try:
            w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
            svg = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">']
            for item in self.canvas.find_all():
                ctype = self.canvas.type(item)
                opts = self.canvas.itemconfig(item)
                fill = opts.get("fill", ("", "", "", "", "none"))[-1]
                outline = opts.get("outline", ("", "", "", "", "none"))[-1]
                width = opts.get("width", (1,))[-1]
                coords = self.canvas.coords(item)
                if ctype == "line":
                    svg.append(f'<line x1="{coords[0]}" y1="{coords[1]}" x2="{coords[2]}" y2="{coords[3]}" stroke="{outline or fill}" stroke-width="{width}"/>')
                elif ctype == "rectangle":
                    svg.append(f'<rect x="{min(coords[0], coords[2])}" y="{min(coords[1], coords[3])}" width="{abs(coords[2]-coords[0])}" height="{abs(coords[3]-coords[1])}" fill="{fill}" stroke="{outline}" stroke-width="{width}"/>')
                elif ctype == "oval":
                    svg.append(f'<ellipse cx="{(coords[0]+coords[2])/2}" cy="{(coords[1]+coords[3])/2}" rx="{abs(coords[2]-coords[0])/2}" ry="{abs(coords[3]-coords[1])/2}" fill="{fill}" stroke="{outline}" stroke-width="{width}"/>')
                elif ctype == "polygon" or (ctype == "line" and len(coords) > 4):
                    pts = " ".join(f"{x},{y}" for x, y in zip(coords[::2], coords[1::2]))
                    svg.append(f'<polyline points="{pts}" fill="none" stroke="{outline or fill}" stroke-width="{width}" stroke-linecap="round"/>')
            svg.append("</svg>")
            with open(path, "w", encoding="utf-8") as f:
                f.write("\n".join(svg))
            messagebox.showinfo("Успех", f"SVG сохранён в {path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Экспорт SVG: {e}")

    # ==================== МЕДИА ПЛЕЕР ====================
    def setup_media_player(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="🎬 Плеер")

        main_split = ttk.Frame(frame)
        main_split.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        left_panel = ttk.Frame(main_split)
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        right_panel = ttk.Frame(main_split)
        right_panel.pack(side=tk.RIGHT, fill=tk.Y, padx=(5, 0))

        self.video_frame = ttk.Frame(left_panel, relief="solid", borderwidth=2)
        self.video_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 5))
        self.video_label = tk.Label(self.video_frame, bg="black")
        self.video_label.pack(fill=tk.BOTH, expand=True)

        ctrl_frame = ttk.Frame(left_panel)
        ctrl_frame.pack(fill=tk.X, pady=5)

        self.file_label = ttk.Label(ctrl_frame, text="Файл не выбран", wraplength=500)
        self.file_label.pack(fill=tk.X, pady=(0, 5))
        self.audio_status = ttk.Label(ctrl_frame, text="", foreground="blue")
        self.audio_status.pack(pady=(0, 5))

        btn_frame = ttk.Frame(ctrl_frame)
        btn_frame.pack(fill=tk.X, pady=5)
        ttk.Button(btn_frame, text="⏮ Пред.", command=self.prev_track).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="▶️ Играть", command=self.play_media).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="⏸ Пауза", command=self.pause_media).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="⏹ Стоп", command=self.stop_media).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="⏭ След.", command=self.next_track).pack(side=tk.LEFT, padx=2)

        self.progress = ttk.Scale(ctrl_frame, from_=0, to=100, orient=tk.HORIZONTAL, state="disabled")
        self.progress.pack(fill=tk.X, pady=5)
        self.progress.bind("<ButtonRelease-1>", self.seek_media)

        vol_frame = ttk.Frame(ctrl_frame)
        vol_frame.pack(fill=tk.X, pady=5)
        ttk.Label(vol_frame, text="🔊").pack(side=tk.LEFT)
        self.volume_label = ttk.Label(vol_frame, text="70%")
        self.volume_label.pack(side=tk.RIGHT)
        self.volume_scale = ttk.Scale(vol_frame, from_=0, to=1, orient=tk.HORIZONTAL, command=self.set_volume)
        self.volume_scale.set(0.7)
        self.volume_scale.pack(side=tk.LEFT, fill=tk.X, expand=True)

        self.time_label = ttk.Label(ctrl_frame, text="00:00 / 00:00")
        self.time_label.pack(pady=2)

        self.playlist_listbox = tk.Listbox(right_panel, bg="#f5f5f5", fg="#000000", selectbackground="#0078d7", selectforeground="#fff", font=("Arial", 10))
        self.playlist_listbox.pack(fill=tk.BOTH, expand=True, pady=(0, 5))
        self.playlist_listbox.bind("<<ListboxSelect>>", self.play_selected)

        pl_toolbar = ttk.Frame(right_panel)
        pl_toolbar.pack(fill=tk.X)
        ttk.Button(pl_toolbar, text="+ Добавить", command=self.add_to_playlist).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(pl_toolbar, text="- Удалить", command=self.remove_from_playlist).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(pl_toolbar, text="Очистить", command=self.clear_playlist).pack(side=tk.LEFT, fill=tk.X, expand=True)

        self.playlist = []
        self.current_playlist_index = -1
        self.media_file = None
        self.is_video = False
        self.is_playing = False
        self.video_cap = None
        self.media_length = 0.0
        self.temp_audio_path = None

        self.update_progress()

    def set_volume(self, val):
        if hasattr(self, 'volume_label'):
            vol = float(val)
            pygame.mixer.music.set_volume(vol)
            self.volume_label.config(text=f"{int(vol*100)}%")

    def seek_media(self, event=None):
        if self.media_length > 0 and self.is_playing:
            val = self.progress.get()
            pos_sec = (val / 100.0) * self.media_length
            try:
                pygame.mixer.music.set_pos(pos_sec)
            except Exception:
                pygame.mixer.music.rewind()
                pygame.mixer.music.play(start=pos_sec)
            if self.is_video:
                self.video_cap.set(cv2.CAP_PROP_POS_MSEC, pos_sec * 1000)

    def add_to_playlist(self):
        files = filedialog.askopenfilenames(filetypes=[("Media", "*.mp3 *.wav *.ogg *.mp4 *.avi *.mkv *.mov")])
        for f in files:
            if f not in self.playlist:
                self.playlist.append(f)
                self.playlist_listbox.insert(tk.END, os.path.basename(f))

    def remove_from_playlist(self):
        idx = self.playlist_listbox.curselection()
        if idx:
            self.playlist.pop(idx[0])
            self.playlist_listbox.delete(idx[0])

    def clear_playlist(self):
        self.playlist.clear()
        self.playlist_listbox.delete(0, tk.END)
        self.stop_media()

    def play_selected(self, event=None):
        idx = self.playlist_listbox.curselection()
        if idx:
            self.current_playlist_index = idx[0]
            self.load_media(self.playlist[idx[0]])

    def load_media(self, path=None):
        if path is None: return
        self.stop_media()
        self.media_file = path
        self.file_label.config(text=f"📁 {os.path.basename(path)}")
        self.progress.config(state="normal")
        self.progress.set(0)
        self.time_label.config(text="00:00 / 00:00")

        ext = os.path.splitext(path)[1].lower()
        try:
            if ext in {".mp3", ".wav", ".ogg"}:
                self.is_video = False
                self._load_audio(path)
                temp = pygame.mixer.Sound(path)
                self.media_length = temp.get_length()
                self.video_frame.pack_forget()
            elif ext in {".mp4", ".avi", ".mkv", ".mov"}:
                self.is_video = True
                self._load_audio(path)
                if self.video_cap: self.video_cap.release()
                self.video_cap = cv2.VideoCapture(path)
                fps = self.video_cap.get(cv2.CAP_PROP_FPS)
                frames = self.video_cap.get(cv2.CAP_PROP_FRAME_COUNT)
                self.media_length = frames / fps if fps > 0 else 10
                self.video_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 5))
                self._show_first_frame()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить: {e}")

    def _has_ffmpeg(self):
        try:
            subprocess.run(["ffmpeg", "-version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
            return True
        except: return False

    def _load_audio(self, path):
        try:
            pygame.mixer.music.load(path)
            self.audio_status.config(text="✅ Встроенный кодек")
        except pygame.error:
            if self._has_ffmpeg():
                self.temp_audio_path = tempfile.mktemp(suffix=".wav")
                subprocess.run(["ffmpeg", "-i", path, "-vn", "-acodec", "pcm_s16le", "-y", self.temp_audio_path],
                               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
                pygame.mixer.music.load(self.temp_audio_path)
                self.audio_status.config(text="✅ Извлечено через ffmpeg")
            else:
                self.audio_status.config(text="⚠️ ffmpeg не найден")

    def _show_first_frame(self):
        ret, frame = self.video_cap.read()
        if ret:
            self.video_cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
            img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            img.thumbnail((700, 400))
            self.tk_img = ImageTk.PhotoImage(image=img)
            self.video_label.config(image=self.tk_img)

    def play_media(self):
        if not self.media_file or self.is_playing: return
        self.is_playing = True
        pygame.mixer.music.play()
        if self.is_video: self._play_video_loop()

    def pause_media(self):
        if self.is_playing:
            pygame.mixer.music.pause()
            self.is_playing = False

    def stop_media(self):
        self.is_playing = False
        pygame.mixer.music.stop()
        if self.is_video and self.video_cap:
            self.video_cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
            self._show_first_frame()
        self.progress.set(0)
        self.time_label.config(text="00:00 / 00:00")

    def next_track(self):
        idx = self.current_playlist_index + 1
        if 0 <= idx < len(self.playlist):
            self.current_playlist_index = idx
            self.playlist_listbox.selection_clear(0, tk.END)
            self.playlist_listbox.selection_set(idx)
            self.load_media(self.playlist[idx])

    def prev_track(self):
        idx = self.current_playlist_index - 1
        if 0 <= idx < len(self.playlist):
            self.current_playlist_index = idx
            self.playlist_listbox.selection_clear(0, tk.END)
            self.playlist_listbox.selection_set(idx)
            self.load_media(self.playlist[idx])

    def _play_video_loop(self):
        if not self.is_playing or not self.is_video or not self.video_cap: return
        ret, frame = self.video_cap.read()
        if ret:
            img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            img.thumbnail((700, 400))
            self.tk_img = ImageTk.PhotoImage(image=img)
            self.video_label.config(image=self.tk_img)
            fps = self.video_cap.get(cv2.CAP_PROP_FPS)
            delay = int(1000 / max(fps, 1))
            self.root.after(delay, self._play_video_loop)
        else:
            self.next_track()

    def update_progress(self):
        # ✅ Безопасный выход при закрытии приложения
        if not self.is_running: return
        if self.is_playing and self.media_length > 0:
            pos = pygame.mixer.music.get_pos() / 1000.0
            if pos < 0: pos = 0
            self.progress.set(min((pos / self.media_length) * 100, 100))
            cur_min, cur_sec = divmod(int(pos), 60)
            tot_min, tot_sec = divmod(int(self.media_length), 60)
            self.time_label.config(text=f"{cur_min:02d}:{cur_sec:02d} / {tot_min:02d}:{tot_sec:02d}")
        self.root.after(100, self.update_progress)

    # ==================== НАСТРОЙКИ ====================
    def setup_settings(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="⚙️ Настройки")

        doc_frame = ttk.LabelFrame(frame, text="📖 Документация проекта", padding=10)
        doc_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # ✅ Исправлено: родитель Scrollbar теперь doc_frame, а не несуществующий doc_text
        self.doc_text = tk.Text(doc_frame, wrap=tk.WORD, font=("Consolas", 10), state=tk.DISABLED)
        doc_scroll = ttk.Scrollbar(doc_frame, command=self.doc_text.yview)
        self.doc_text.config(yscrollcommand=doc_scroll.set, state=tk.NORMAL)
        
        self.doc_text.insert(tk.END, self.get_documentation())
        self.doc_text.config(state=tk.DISABLED)
        
        self.doc_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        doc_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        dev_frame = ttk.LabelFrame(frame, text="👨‍💻 Группа разработки", padding=10)
        dev_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(dev_frame, text="Команда: sahariadev", font=("Arial", 12, "bold")).pack(anchor=tk.W, pady=2)
        ttk.Label(dev_frame, text="Версия: v6.1 Stable", font=("Arial", 10)).pack(anchor=tk.W, pady=2)
        ttk.Label(dev_frame, text="Лицензия: MIT (Open Source)", font=("Arial", 10)).pack(anchor=tk.W, pady=2)
        ttk.Label(dev_frame, text="Описание: Мультиплатформенный инструмент на Python/Tkinter, объединяющий текстовый редактор, графический холст и медиаплеер.", font=("Arial", 10)).pack(anchor=tk.W, pady=(5, 0))

        ttk.Button(frame, text="🌗 Переключить тему", command=self.toggle_theme).pack(pady=15)

    def get_documentation(self):
        return """📌 ГОРЯЧИЕ КЛАВИШИ:
• Ctrl+N / Ctrl+O / Ctrl+S  - Новый / Открыть / Сохранить файл
• Ctrl+A / Ctrl+C / Ctrl+X / Ctrl+V - Выделить всё / Копировать / Вырезать / Вставить
• Ctrl+Z / Ctrl+Y           - Отмена / Повтор действия
• F1 / F2                   - Тёмная / Светлая тема
• Space                     - Play / Pause (в медиаплеере)

📌 ТЕКСТОВЫЙ РЕДАКТОР:
• Поддержка .txt, .md, .log, .docx
• Встраиваемые таблицы и настраиваемые шрифты
• Подсветка синтаксиса Python (ключевые слова, строки, комментарии, числа)

📌 PAINT:
• 5 инструментов: карандаш, линия, прямоугольник, овал, ластик
• Предпросмотр фигур при рисовании
• Экспорт в растр (PNG/JPG) и вектор (SVG)

📌 МЕДИА ПЛЕЕР:
• Аудио: MP3, WAV, OGG
• Видео: MP4, AVI, MKV, MOV (требуется ffmpeg для звука)
• Плейлист с авто-переключением, перемотка кликом, регулятор громкости

📌 ЗАВИСИМОСТИ:
• pygame, opencv-python, Pillow, python-docx
• ffmpeg (в PATH системы) для извлечения аудио из видео

✅ Приложение полностью автономно и не требует установки дополнительных драйверов."""

    # ==================== ГОРЯЧИЕ КЛАВИШИ ====================
    def setup_global_hotkeys(self):
        self.root.bind("<Control-n>", lambda e: self.new_file())
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-s>", lambda e: self.save_file())
        
        self.text.bind("<Control-a>", self.select_all)
        self.text.bind("<Control-z>", lambda e: self.text.edit_undo())
        self.text.bind("<Control-y>", lambda e: self.text.edit_redo())
        self.text.bind("<Control-c>", lambda e: self.text.event_generate("<<Copy>>"))
        self.text.bind("<Control-x>", lambda e: self.text.event_generate("<<Cut>>"))
        self.text.bind("<Control-v>", lambda e: self.text.event_generate("<<Paste>>"))
        
        self.root.bind("<F1>", lambda e: self.apply_theme("dark"))
        self.root.bind("<F2>", lambda e: self.apply_theme("light"))
        self.root.bind("<space>", lambda e: (self.play_media() if not self.is_playing else self.pause_media()) or "break")

    def select_all(self, event=None):
        self.text.tag_add("sel", "1.0", "end")
        return "break"

    def on_close(self):
        self.is_running = False  # Останавливаем все after-циклы
        if self.temp_audio_path and os.path.exists(self.temp_audio_path):
            os.remove(self.temp_audio_path)
        if self.video_cap: self.video_cap.release()
        pygame.mixer.quit()
        self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    try:
        app = MultiToolApp(root)
        root.mainloop()
    except Exception as e:
        print(f"Критическая ошибка: {e}")
        input("Enter для выхода...")
