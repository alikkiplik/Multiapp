import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import os, re

class TextEditorModule(ttk.Frame):
    def __init__(self, parent, theme_mgr):
        super().__init__(parent)
        self.theme_mgr = theme_mgr
        self.current_file = None
        self.syntax_timer = None
        self.highlight_syntax = tk.BooleanVar(value=False)
        self._setup_ui()
        self._setup_syntax_tags()
        self._bind_hotkeys()

    def _setup_ui(self):
        toolbar = ttk.Frame(self)
        toolbar.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(toolbar, text="📄 Новый", command=self.new_file).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="📂 Открыть", command=self.open_file).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="💾 Сохранить", command=self.save_file).pack(side=tk.LEFT, padx=2)
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)
        
        self.fonts = [
            "Arial", "Calibri", "Consolas", "Courier New", "DejaVu Sans Mono",
            "Georgia", "Helvetica", "JetBrains Mono", "Noto Sans", "Roboto",
            "Segoe UI", "Tahoma", "Times New Roman", "Ubuntu Mono", "Verdana"
        ]
        self.sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 28, 32]
        self.font_combo = ttk.Combobox(toolbar, values=self.fonts, width=14, state="readonly")
        self.font_combo.current(0)
        self.font_combo.pack(side=tk.LEFT, padx=2)
        self.size_combo = ttk.Combobox(toolbar, values=self.sizes, width=3, state="readonly")
        self.size_combo.current(4)
        self.size_combo.pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="🔤 Шрифт", command=self.apply_font).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="📊 Таблица", command=self.insert_table).pack(side=tk.LEFT, padx=2)
        ttk.Checkbutton(toolbar, text="🎨 Подсветка", variable=self.highlight_syntax, command=self.toggle_highlighting).pack(side=tk.LEFT, padx=10)
        
        self.text = tk.Text(self, wrap=tk.WORD, font=("Consolas", 12), undo=True)
        self.text.pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5))
        
        self.status_var = tk.StringVar(value="Готов к работе | Ctrl+N/O/S | F1/F2 Тема")
        ttk.Label(self, textvariable=self.status_var, anchor=tk.W).pack(fill=tk.X, padx=5)
        self.text.bind("<KeyRelease>", self.schedule_highlight)

    def _setup_syntax_tags(self):
        self.text.tag_configure("keyword", foreground="#569cd6", font=("Consolas", 12, "bold"))
        self.text.tag_configure("string", foreground="#ce9178")
        self.text.tag_configure("comment", foreground="#6a9955", font=("Consolas", 12, "italic"))
        self.text.tag_configure("number", foreground="#b5cea8")

    def _bind_hotkeys(self):
        # ✅ Привязка к классу Text гарантирует работу даже если фокус внутри виджета
        self.text.bind_class("Text", "<Control-a>", lambda e: (self.text.tag_add("sel", "1.0", "end"), "break"), add="+")
        self.text.bind_class("Text", "<Control-z>", lambda e: (self.text.edit_undo(), "break"), add="+")
        self.text.bind_class("Text", "<Control-y>", lambda e: (self.text.edit_redo(), "break"), add="+")

    def apply_colors(self, colors):
        self.text.config(bg=colors["text_bg"], fg=colors["fg"], 
                         insertbackground=colors["fg"], selectbackground=colors["accent"])

    def schedule_highlight(self, event=None):
        if self.syntax_timer: self.master.after_cancel(self.syntax_timer)
        if self.highlight_syntax.get(): self.syntax_timer = self.master.after(300, self.highlight_text)

    def highlight_text(self):
        if not self.highlight_syntax.get(): return
        txt = self.text.get("1.0", "end-1c")
        for tag in ("keyword", "string", "comment", "number"): self.text.tag_remove(tag, "1.0", tk.END)
        keywords = ["def", "class", "import", "from", "if", "else", "for", "while", "return", "try", "except", "with", "as", "True", "False", "None", "and", "or", "not"]
        for line_num, line in enumerate(txt.split("\n"), 1):
            for m in re.finditer(r"(#.*)", line): self.text.tag_add("comment", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")
            for m in re.finditer(r"(['\"].*?['\"])", line): self.text.tag_add("string", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")
            for m in re.finditer(r"\b(\d+\.?\d*)\b", line): self.text.tag_add("number", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")
            for kw in keywords:
                for m in re.finditer(rf"\b{kw}\b", line): self.text.tag_add("keyword", f"{line_num}.{m.start()}", f"{line_num}.{m.end()}")

    def toggle_highlighting(self):
        if not self.highlight_syntax.get():
            for tag in ("keyword", "string", "comment", "number"): self.text.tag_remove(tag, "1.0", tk.END)

    def new_file(self): self.text.delete("1.0", tk.END); self.current_file = None
    def open_file(self):
        path = filedialog.askopenfilename(filetypes=[("Text/Code/Word", "*.txt *.md *.log *.py *.js *.html *.css *.docx")])
        if not path: return
        try:
            self.text.delete("1.0", tk.END)
            content = ""
            for enc in ["utf-8", "utf-8-sig", "cp1251"]:
                try:
                    with open(path, "r", encoding=enc) as f: content = f.read(); break
                except UnicodeDecodeError: continue
            if path.endswith(".docx"):
                from docx import Document
                doc = Document(path)
                content = "\n".join([p.text for p in doc.paragraphs])
            self.text.insert(tk.END, content)
            self.current_file = path
            self.status_var.set(f"Открыт: {os.path.basename(path)}")
        except Exception as e: messagebox.showerror("Ошибка", str(e))

    def save_file(self):
        if self.current_file: path = self.current_file
        else: path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text", "*.txt"), ("Word", "*.docx")])
        if not path: return
        try:
            if path.endswith(".docx"):
                from docx import Document
                doc = Document()
                for line in self.text.get("1.0", "end-1c").split("\n"): doc.add_paragraph(line)
                doc.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f: f.write(self.text.get("1.0", tk.END))
            self.current_file = path
            self.status_var.set(f"Сохранено: {os.path.basename(path)}")
        except Exception as e: messagebox.showerror("Ошибка", str(e))

    def apply_font(self):
        sel = self.text.tag_ranges("sel")
        if not sel: return messagebox.showinfo("Инфо", "Выделите текст.")
        self.text.tag_configure("custom_font", font=(self.font_combo.get(), int(self.size_combo.get())))
        self.text.tag_add("custom_font", sel[0], sel[1])

    def insert_table(self):
        rows = simpledialog.askinteger("Таблица", "Строки:", minvalue=1, maxvalue=20, initialvalue=3)
        cols = simpledialog.askinteger("Таблица", "Столбцы:", minvalue=1, maxvalue=20, initialvalue=3)
        if rows and cols:
            table_frame = ttk.Frame(self.text, relief=tk.SOLID, borderwidth=1)
            for r in range(rows):
                for c in range(cols): ttk.Entry(table_frame, width=12).grid(row=r, column=c, padx=1, pady=1)
            self.text.window_create(tk.INSERT, window=table_frame)
            self.text.insert(tk.INSERT, "\n\n")
